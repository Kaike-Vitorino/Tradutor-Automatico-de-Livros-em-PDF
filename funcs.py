from PyPDF2 import *
from deep_translator import GoogleTranslator
import spacy
import PyPDF2
from docx import Document
from docx.shared import Pt
import re
from docx2pdf import convert
import os

def add_pages_anteriores_pdf(arquivo_pdf, numero_pagina, arquivo_pdf_pronto):
    # Diminuindo em 1 para que a primeira pagina nao seja copiada
    numero_pagina_antes = numero_pagina - 1

    # Abrindo o arquivo PDF original
    with open(arquivo_pdf, 'rb') as arquivo_original:
        # Criando um objeto PDFReader
        leitor_pdf = PyPDF2.PdfReader(arquivo_original)

        # Criando um objeto PDFWriter
        escritor_pdf = PyPDF2.PdfWriter()

        # Adicionando as páginas do PDF original ao PDF traduzido
        for num_pagina in range(numero_pagina_antes):
            pagina = leitor_pdf.pages[num_pagina]
            escritor_pdf.add_page(pagina)

        # Modificando o PDF traduzido
        with open(arquivo_pdf_pronto, 'wb') as novo_arquivo:
            escritor_pdf.write(novo_arquivo)

def extrair_pagina(arquivo_pdf, numero_pagina):
    with open(arquivo_pdf, 'rb') as f:
        leitor_pdf = PdfReader(f)
        escritor_pdf = PdfWriter()

        escritor_pdf.add_page(leitor_pdf.pages[numero_pagina - 1])

        with open('pagina.pdf', 'wb') as pdf_saida:
            escritor_pdf.write(pdf_saida)


def converter_pagina_em_texto(pagina):
    pdf_leitor = PdfReader(pagina)

    pagina_conteudo = {}

    for indx, pdf_pag in enumerate(pdf_leitor.pages):
        pagina_conteudo[indx + 1] = pdf_pag.extract_text()

    return pagina_conteudo

def texto_traduzir(conteudo_pagina, idioma_requerido):
    # Traduzir o conteúdo da página para o idioma requerido usando o DeepL
    texto_traduzido = GoogleTranslator(source='auto', target=idioma_requerido).translate(text=conteudo_pagina)

    return texto_traduzido

def traduzir_com_repeticao(conteudo_pagina, idioma_requerido, tentativas=3):
    for _ in range(tentativas):
        texto_traduzido = texto_traduzir(conteudo_pagina, idioma_requerido)
        # Verificar se todas as palavras foram traduzidas
        palavras_originais = conteudo_pagina.split()
        palavras_traduzidas = texto_traduzido.split()

        if all(palavra_traduzida != palavra_original for palavra_original, palavra_traduzida in zip(palavras_originais, palavras_traduzidas)):
            return texto_traduzido

        else:
            return texto_traduzido

def limpar_texto(texto):
    # Remover caracteres indesejados
    texto = texto.replace(r'\n', '\n')  # Substituir quebras de linha por espaços
    texto = texto.replace('{', '')  # Remover chaves abertas
    texto = texto.replace('}', '')  # Remover chaves fechadas
    texto = texto.replace("'", '')  # Remover aspas simples
    texto = texto.replace("Goldenagato | mp4directs.com", '')  # Remover propaganda
    texto = re.sub(r'^\d+: ', '', texto, flags=re.MULTILINE)  # Remover "1: " avulsos no início de uma linha
    return texto

def preservar_nomes_proprios(texto, idioma='pt_core_news_sm'):
    # Carregue o modelo de linguagem spaCy
    nlp = spacy.load(idioma)

    # Processar o texto
    doc = nlp(texto)

    # Dicionário para mapear cada nome próprio para um marcador único
    nomes_proprios = {}

    # Substituir nomes próprios por um marcador único
    for ent in doc.ents:
        if ent.label_ == 'PER':
            marcador = '<NOMEPROPRIO{}>'.format(len(nomes_proprios))
            texto = texto.replace(ent.text, marcador)
            nomes_proprios[marcador] = ent.text

    return texto, nomes_proprios

# Função para criar um novo documento Word
def criar_arquivo_word(nome_arquivo):
    doc = Document()
    doc.save(nome_arquivo)

# Função para adicionar uma página traduzida ao documento Word
def add_pags_no_word(arquivo_word_pronto, pagina_traduzida_estilizada):
    # Abre ou cria o documento Word
    try:
        doc = Document(arquivo_word_pronto)
    except FileNotFoundError:
        criar_arquivo_word(arquivo_word_pronto)
        doc = Document(arquivo_word_pronto)

    # Adiciona a nova página traduzida como um novo parágrafo
    paragrafo = doc.add_paragraph()

    # Cria um novo run com o texto traduzido e define o tamanho da fonte para 13
    run = paragrafo.add_run(pagina_traduzida_estilizada)
    run.font.size = Pt(13)

    # Substitui a contagem antiga
    paragrafo.text = re.sub(r'Página \d+', '', paragrafo.text, flags=re.IGNORECASE)

    # Altera o tamanho da fonte de todo o texto já adicionado para 13
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(13)

    # Salva o documento Word
    doc.save(arquivo_word_pronto)

#COnvertendo word para pdf
def converter_word_para_pdf(arquivo_word, arquivo_pdf_temporario):
    convert(arquivo_word, arquivo_pdf_temporario)

# Juntando o pdf convertido com o pdf com as imagens inicias do livro
def juntar_pdfs(arquivo_pdf_pronto, arquivo_pdf_temporario):
    with open(arquivo_pdf_pronto, 'rb') as pdf_existente, open(arquivo_pdf_temporario, 'rb') as pdf_temporario:
        pdf_existente_reader = PdfReader(pdf_existente)
        pdf_temporario_reader = PdfReader(pdf_temporario)

        escreve_pdf_final = PdfWriter()

        # Adicionar páginas do PDF existente
        for num_pagina in range(len(pdf_existente_reader.pages)):
            pagina_final = pdf_existente_reader.pages[num_pagina]
            escreve_pdf_final.add_page(pagina_final)

        # Adicionar páginas do PDF temporário
        for num_pagina in range(len(pdf_temporario_reader.pages)):
            pagina_temporaria = pdf_temporario_reader.pages[num_pagina]
            escreve_pdf_final.add_page(pagina_temporaria)

        # Salvar o PDF final
        with open(arquivo_pdf_pronto, 'wb') as pdf_final:
            escreve_pdf_final.write(pdf_final)

# FUnc que vai puxar as duas funcoes acima
def processar_arquivos(arquivo_word, arquivo_pdf_existente):
    # Criar PDF temporário
    arquivo_pdf_temporario = 'temporario.pdf'
    converter_word_para_pdf(arquivo_word, arquivo_pdf_temporario)

    # Juntar arquivos PDF
    juntar_pdfs(arquivo_pdf_existente, arquivo_pdf_temporario)

    # Remover arquivos temporários (opcional)
    os.remove(arquivo_word)
    os.remove(arquivo_pdf_temporario)

    print(f"Processo concluído.")
